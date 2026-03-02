import sys
import asyncio
import traceback
import io
import csv
import os
import json
import zipfile
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles

from docx import Document
import httpx

app = FastAPI(title="DOCX Summarizer")
print(f"[startup] encoding: {sys.getdefaultencoding()}, stdout: {sys.stdout.encoding}", flush=True)

app.mount("/static", StaticFiles(directory="static"), name="static")

QWEN_URL = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions"

SUMMARIZE_PROMPT = (
    "Сделай суммаризацию текста на русском языке. "
    "Объём — не более 1000 символов. "
    "Структурируй по логическим разделам документа, используя короткие подзаголовки. "
    "Отвечай только суммаризацией, без вводных фраз."
)

MAX_FILES = 20


async def call_qwen(text: str) -> str:
    """Call Qwen3-Omni-Flash API via direct httpx (streaming SSE), UTF-8 safe."""
    api_key = os.environ.get("QWEN_API_KEY", "")
    payload = {
        "model": "qwen3-omni-flash",
        "messages": [
            {"role": "system", "content": SUMMARIZE_PROMPT},
            {"role": "user", "content": text},
        ],
        "stream": True,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
    }
    # json= в httpx всегда кодирует как UTF-8, независимо от системной локали
    payload_bytes = json.dumps(payload, ensure_ascii=False).encode("utf-8")

    parts = []
    async with httpx.AsyncClient(timeout=120.0) as client:
        async with client.stream(
            "POST", QWEN_URL, content=payload_bytes, headers=headers
        ) as response:
            response.raise_for_status()
            async for line in response.aiter_lines():
                line = line.strip()
                if not line or line == "data: [DONE]":
                    continue
                if line.startswith("data: "):
                    try:
                        data = json.loads(line[6:])
                        content = (data["choices"][0]["delta"] or {}).get("content") or ""
                        parts.append(content)
                    except (json.JSONDecodeError, KeyError, IndexError):
                        pass

    return "".join(parts).strip()


@app.get("/")
async def serve_index():
    return FileResponse("static/index.html")


@app.post("/convert")
async def convert_docx(files: list[UploadFile] = File(...)):
    """Convert DOCX files to TXT and return them as a ZIP archive."""
    if not files:
        raise HTTPException(status_code=400, detail="Файлы не выбраны")

    files = files[:MAX_FILES]

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for upload in files:
            try:
                content = await upload.read()
                doc = Document(io.BytesIO(content))
                text = "\n".join(para.text for para in doc.paragraphs)
                txt_bytes = ("\uFEFF" + text).encode("utf-8")
                txt_name = Path(upload.filename).stem + ".txt"
                zf.writestr(txt_name, txt_bytes)
            except Exception as e:
                txt_name = Path(upload.filename).stem + "_ERROR.txt"
                zf.writestr(txt_name, f"Ошибка конвертации: {e}".encode("utf-8"))

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=converted_txts.zip"},
    )


@app.post("/summarize")
async def summarize_txt(files: list[UploadFile] = File(...)):
    """Summarize DOCX files via Qwen3-Omni-Flash and return a CSV."""
    if not files:
        raise HTTPException(status_code=400, detail="Файлы не выбраны")

    files = files[:MAX_FILES]

    async def process_one(upload: UploadFile):
        name = Path(upload.filename).stem
        try:
            raw = await upload.read()
            doc = Document(io.BytesIO(raw))
            text = "\n".join(para.text for para in doc.paragraphs).strip()
            summary = await call_qwen(text)
        except Exception as e:
            tb = traceback.format_exc()
            print(f"[summarize ERROR for '{name}'] {tb}", flush=True)
            summary = f"Ошибка обработки: {e}"
        return {"Имя файла": name, "Суммаризация": summary}

    # Параллельная обработка всех файлов — вместо последовательной
    rows = list(await asyncio.gather(*[process_one(u) for u in files]))

    # UTF-8 with BOM + semicolons — корректно открывается в Excel с кириллицей
    csv_buffer = io.StringIO()
    writer = csv.DictWriter(
        csv_buffer,
        fieldnames=["Имя файла", "Суммаризация"],
        delimiter=";",
        quoting=csv.QUOTE_ALL,
    )
    writer.writeheader()
    writer.writerows(rows)

    csv_bytes = ("\uFEFF" + csv_buffer.getvalue()).encode("utf-8")

    return StreamingResponse(
        io.BytesIO(csv_bytes),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=summaries.csv"},
    )
